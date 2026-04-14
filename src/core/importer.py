from __future__ import annotations

import csv
import html
import re
from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from openpyxl import load_workbook

from src.core.generic_import import items_from_rows, items_from_text_blocks
from src.core.gbt37668 import (
    GBT37668_HEADERS,
    GBT37668_TEMPLATE_NAME,
    item_from_gbt37668_mapping,
)
from src.core.project_manager import ProjectData, ProjectItem, ProjectMeta
from src.core.structured_import import TableSchema, parse_tables_with_schema

MARKDOWN_NEWLINE_TOKEN = "<!--A11Y-MD-NL-->"

STATUS_MAP = {
    "pending": "pending",
    "待处理": "pending",
    "passed": "passed",
    "通过": "passed",
    "failed": "failed",
    "失败": "failed",
    "not_applicable": "not_applicable",
    "不适用": "not_applicable",
}

PRIORITY_MAP = {
    "low": "low",
    "低": "low",
    "medium": "medium",
    "中": "medium",
    "high": "high",
    "高": "high",
}

DEFAULT_TEMPLATES = {
    "html": "HTML 导入",
    "word": "Word 导入",
    "excel": "Excel 导入",
    "jira_csv": "CSV 导入",
    "markdown": "Markdown 导入",
}

GENERIC_TEMPLATES = {
    "html": "HTML 文档导入",
    "word": "Word 文档导入",
    "excel": "Excel 文档导入",
    "jira_csv": "CSV 文档导入",
    "markdown": "Markdown 文档导入",
}


@dataclass(slots=True)
class ImportedProject:
    data: ProjectData
    source_format: str


@dataclass(slots=True)
class DocumentContent:
    text_blocks: list[str]
    tables: list[list[list[str]]]


@dataclass(frozen=True, slots=True)
class StructuredResult:
    items: list[ProjectItem]
    template_name: str


class _HtmlImportParser(HTMLParser):
    _BLOCK_TAGS = {"div", "p", "li", "h1", "h2", "h3", "h4", "h5", "h6"}

    def __init__(self) -> None:
        super().__init__()
        self.in_text_block = False
        self.in_table = False
        self.current_table: list[list[str]] = []
        self.current_cell: list[str] = []
        self.current_row: list[str] = []
        self.text_buffer: list[str] = []
        self.text_blocks: list[str] = []
        self.tables: list[list[list[str]]] = []
        self._cell_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in self._BLOCK_TAGS and not self._cell_depth:
            self.in_text_block = True
            self.text_buffer = []
        elif tag == "table":
            self.in_table = True
            self.current_table = []
        elif tag == "tr" and self.in_table:
            self.current_row = []
        elif tag in {"td", "th"} and self.in_table:
            self._cell_depth += 1
            if self._cell_depth == 1:
                self.current_cell = []
        elif tag == "br" and self._cell_depth:
            self.current_cell.append("\n")
        elif tag == "br" and self.in_text_block:
            self.text_buffer.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._BLOCK_TAGS and self._cell_depth:
            self._append_cell_break()
        elif tag in self._BLOCK_TAGS:
            text = "".join(self.text_buffer).strip()
            if text:
                self.text_blocks.append(text)
            self.in_text_block = False
            self.text_buffer = []
        elif tag in {"td", "th"} and self._cell_depth:
            self._cell_depth -= 1
            if self._cell_depth == 0:
                self.current_row.append("".join(self.current_cell).strip())
                self.current_cell = []
        elif tag == "tr" and self.in_table:
            if self.current_row and any(cell.strip() for cell in self.current_row):
                self.current_table.append(self.current_row)
            self.current_row = []
        elif tag == "table" and self.in_table:
            if self.current_table:
                self.tables.append(self.current_table)
            self.current_table = []
            self.in_table = False

    def handle_data(self, data: str) -> None:
        text = html.unescape(data)
        if self.in_text_block:
            self.text_buffer.append(text)
        elif self._cell_depth:
            self.current_cell.append(text)

    def _append_cell_break(self) -> None:
        if self.current_cell and not self.current_cell[-1].endswith("\n"):
            self.current_cell.append("\n")


class ProjectImporter:
    def import_file(self, path: Path, format_name: str) -> ImportedProject:
        handlers = {
            "html": self._import_html,
            "word": self._import_word,
            "excel": self._import_excel,
            "jira_csv": self._import_csv,
            "markdown": self._import_markdown,
        }
        try:
            handler = handlers[format_name]
        except KeyError as exc:
            raise RuntimeError(f"暂不支持的导入格式：{format_name}") from exc
        return handler(path)

    def _import_html(self, path: Path) -> ImportedProject:
        parser = _HtmlImportParser()
        parser.feed(path.read_text(encoding="utf-8"))
        content = DocumentContent(text_blocks=parser.text_blocks, tables=parser.tables)
        return self._build_project(path, "html", content)

    def _import_word(self, path: Path) -> ImportedProject:
        document = Document(path)
        text_blocks = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        tables = [
            [self._row_values(row) for row in table.rows] for table in document.tables
        ]
        content = DocumentContent(text_blocks=text_blocks, tables=tables)
        return self._build_project(path, "word", content)

    def _import_excel(self, path: Path) -> ImportedProject:
        workbook = load_workbook(path, data_only=True)
        text_blocks = [
            sheet.title for sheet in workbook.worksheets if sheet.title.strip()
        ]
        tables = [
            [
                [self._string(value).strip() for value in row]
                for row in sheet.iter_rows(values_only=True)
            ]
            for sheet in workbook.worksheets
        ]
        content = DocumentContent(text_blocks=text_blocks, tables=tables)
        return self._build_project(path, "excel", content)

    def _import_csv(self, path: Path) -> ImportedProject:
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.reader(handle))
        content = DocumentContent(text_blocks=[], tables=[rows] if rows else [])
        return self._build_project(path, "jira_csv", content)

    def _import_markdown(self, path: Path) -> ImportedProject:
        lines = path.read_text(encoding="utf-8").splitlines()
        content = DocumentContent(
            text_blocks=self._markdown_text_blocks(lines),
            tables=self._markdown_tables(lines),
        )
        return self._build_project(path, "markdown", content)

    def _build_project(
        self, path: Path, source_format: str, content: DocumentContent
    ) -> ImportedProject:
        default_template = DEFAULT_TEMPLATES[source_format]
        meta = self._meta_from_text_blocks(content.text_blocks, path, default_template)

        structured = self._parse_structured_tables(content.tables)
        if structured:
            items = structured.items
            if meta.template == default_template:
                meta.template = structured.template_name
        else:
            items = self._generic_items_from_content(content, source_format)
            if items and meta.template == default_template:
                meta.template = GENERIC_TEMPLATES[source_format]

        return ImportedProject(ProjectData(meta=meta, items=items), source_format)

    def _parse_structured_tables(
        self, tables: Sequence[Sequence[Sequence[str]]]
    ) -> StructuredResult | None:
        for schema in self._table_schemas():
            items = parse_tables_with_schema(tables, schema)
            if items:
                return StructuredResult(
                    items=items,
                    template_name=schema.template_name or "结构化导入",
                )
        return None

    def _generic_items_from_content(
        self, content: DocumentContent, source_format: str
    ) -> list[ProjectItem]:
        table_items: list[ProjectItem] = []
        for index, table in enumerate(content.tables, start=1):
            prefix = self._generic_row_prefix(source_format, index)
            table_items.extend(items_from_rows(table, prefix=prefix))
        if table_items:
            return table_items
        return items_from_text_blocks(
            content.text_blocks, prefix=self._generic_text_prefix(source_format)
        )

    def _table_schemas(self) -> list[TableSchema]:
        return [
            TableSchema(
                name="report-check-item",
                required_headers=("ID", "检查项", "状态", "优先级", "描述"),
                row_factory=self._report_item_from_mapping,
                template_name="检查报告导入",
            ),
            TableSchema(
                name="report-content",
                required_headers=("ID", "内容", "状态", "优先级", "描述"),
                row_factory=self._report_item_from_mapping,
                template_name="检查报告导入",
            ),
            TableSchema(
                name="gbt37668",
                required_headers=tuple(GBT37668_HEADERS),
                row_factory=self._gbt37668_item_from_mapping,
                template_name=GBT37668_TEMPLATE_NAME,
            ),
            TableSchema(
                name="jira",
                required_headers=("Summary", "Description", "Priority"),
                row_factory=self._jira_item_from_mapping,
                template_name="Jira CSV 导入",
            ),
        ]

    def _meta_from_text_blocks(
        self, text_blocks: Sequence[str], path: Path, default_template: str
    ) -> ProjectMeta:
        merged = "\n".join(text_blocks)
        return ProjectMeta(
            project_number=self._extract_value(merged, "项目编号"),
            created_time=self._extract_value(merged, "创建时间"),
            scenario=self._extract_value(merged, "场景") or "导入",
            template=self._extract_value(merged, "模板") or default_template,
            project_name=(
                self._extract_value(merged, "工程名称")
                or self._extract_value(merged, "名称")
                or path.stem
            ),
        )

    def _extract_value(self, text: str, label: str) -> str:
        match = re.search(rf"{re.escape(label)}[:：]\s*([^\n]+)", text)
        return match.group(1).strip() if match else ""

    def _report_item_from_mapping(
        self, row: Mapping[str, str], _index: int
    ) -> ProjectItem | None:
        item_id = self._mapped_value(row, "ID")
        content = self._mapped_value(row, "检查项", "内容")
        if not item_id and not content:
            return None
        return ProjectItem(
            id=item_id or "IMPORT",
            content=content or "导入内容",
            status=self._normalize_status(self._mapped_value(row, "状态")),
            priority=self._normalize_priority(self._mapped_value(row, "优先级")),
            description=self._normalize_empty(self._mapped_value(row, "描述")),
            image_path=self._normalize_empty(self._mapped_value(row, "截图", "图片")),
        )

    def _gbt37668_item_from_mapping(
        self, row: Mapping[str, str], index: int
    ) -> ProjectItem | None:
        return item_from_gbt37668_mapping(row, index)

    def _jira_item_from_mapping(
        self, row: Mapping[str, str], index: int
    ) -> ProjectItem | None:
        summary = self._mapped_value(row, "Summary")
        description = self._mapped_value(row, "Description")
        if not summary and not description:
            return None
        item_id, content = self._parse_jira_summary(summary, index)
        return ProjectItem(
            id=item_id,
            content=content,
            status="failed",
            priority=self._normalize_priority(self._mapped_value(row, "Priority")),
            description=description,
        )

    def _parse_jira_summary(self, summary: str, index: int) -> tuple[str, str]:
        cleaned = summary.replace("[A11Y]", "", 1).strip()
        match = re.match(r"([A-Za-z0-9_:-]+)\s+(.*)", cleaned)
        if match:
            return match.group(1), match.group(2).strip()
        return f"IMPORT_{index:03d}", cleaned or f"导入问题 {index}"

    def _mapped_value(self, row: Mapping[str, str], *keys: str) -> str:
        for key in keys:
            value = self._string(row.get(key, "")).strip()
            if value:
                return value
        return ""

    def _markdown_tables(self, lines: Sequence[str]) -> list[list[list[str]]]:
        tables: list[list[list[str]]] = []
        current_table: list[list[str]] = []
        for line in lines:
            stripped = line.strip()
            if not stripped.startswith("|"):
                if current_table:
                    tables.append(current_table)
                    current_table = []
                continue
            if self._is_markdown_separator(stripped):
                continue
            current_table.append(self._split_markdown_row(stripped))
        if current_table:
            tables.append(current_table)
        return tables

    def _markdown_text_blocks(self, lines: Sequence[str]) -> list[str]:
        blocks: list[str] = []
        for line in lines:
            stripped = line.strip()
            if not stripped or stripped.startswith("|"):
                continue
            stripped = re.sub(r"^#{1,6}\s*", "", stripped)
            stripped = re.sub(r"^[-*+]\s+", "", stripped)
            if stripped:
                blocks.append(stripped)
        return blocks

    def _is_markdown_separator(self, line: str) -> bool:
        columns = [column.strip() for column in line.strip("|").split("|")]
        return bool(columns) and all(
            column and set(column) <= {"-", ":"} for column in columns
        )

    def _split_markdown_row(self, line: str) -> list[str]:
        content = line.strip().strip("|")
        if not content:
            return [""]

        columns: list[str] = []
        current: list[str] = []
        escaped = False
        for char in content:
            if escaped:
                if char in {"|", "\\"}:
                    current.append(char)
                else:
                    current.append("\\")
                    current.append(char)
                escaped = False
                continue
            if char == "\\":
                escaped = True
                continue
            if char == "|":
                columns.append(self._normalize_markdown_cell("".join(current)))
                current = []
                continue
            current.append(char)
        if escaped:
            current.append("\\")
        columns.append(self._normalize_markdown_cell("".join(current)))
        return columns

    def _normalize_markdown_cell(self, value: str) -> str:
        return value.replace(MARKDOWN_NEWLINE_TOKEN, "\n").strip()

    def _generic_row_prefix(self, source_format: str, table_index: int) -> str:
        base = {
            "html": "HTMLROW",
            "word": "WORDTBL",
            "excel": "XLSX",
            "jira_csv": "CSV",
            "markdown": "MDTBL",
        }[source_format]
        return f"{base}_{table_index:02d}"

    def _generic_text_prefix(self, source_format: str) -> str:
        return {
            "html": "HTML",
            "word": "WORD",
            "excel": "XLSX",
            "jira_csv": "CSV",
            "markdown": "MD",
        }[source_format]

    def _normalize_status(self, value: str) -> str:
        cleaned = value.strip()
        return STATUS_MAP.get(cleaned.lower(), STATUS_MAP.get(cleaned, "pending"))

    def _normalize_priority(self, value: str) -> str:
        return PRIORITY_MAP.get(
            value.strip().lower(), PRIORITY_MAP.get(value.strip(), "medium")
        )

    def _normalize_empty(self, value: str) -> str:
        cleaned = self._string(value).strip()
        return "" if cleaned in {"", "-"} else cleaned

    def _row_values(self, row) -> list[str]:
        return [self._string(cell.text).strip() for cell in row.cells]

    def _string(self, value) -> str:
        return "" if value is None else str(value)
