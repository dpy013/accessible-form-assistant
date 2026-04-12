from __future__ import annotations

import csv
from dataclasses import asdict
from pathlib import Path

from docx import Document
from jinja2 import Environment, FileSystemLoader, select_autoescape
from openpyxl import Workbook

from src.core.project_manager import ProjectItem, ProjectSession


class ProjectExporter:
    def __init__(self) -> None:
        self.template_dir = Path(__file__).resolve().parent.parent / "templates"
        self.environment = Environment(
            loader=FileSystemLoader(str(self.template_dir)),
            autoescape=select_autoescape(["html"]),
        )
        self.styles = (self.template_dir / "styles.css").read_text(encoding="utf-8")

    def export_html(
        self, session: ProjectSession, destination: Path | None = None
    ) -> Path:
        target = destination or (session.root / "report.html")
        template = self.environment.get_template("html_template.html")
        html = template.render(
            meta=asdict(session.data.meta),
            items=self._active_items(session),
            stats=self._stats(session),
            styles=self.styles,
        )
        target.write_text(html, encoding="utf-8")
        return target

    def export_word(
        self, session: ProjectSession, destination: Path | None = None
    ) -> Path:
        target = destination or (session.root / "report.docx")
        document = Document()
        document.add_heading("信息无障碍测试报告", level=0)
        title_parts = [
            f"项目编号：{session.data.meta.project_number}",
            f"场景：{session.data.meta.scenario}",
            f"模板：{session.data.meta.template}",
        ]
        if session.data.meta.project_name:
            title_parts.insert(1, f"名称：{session.data.meta.project_name}")
        document.add_paragraph("    ".join(title_parts))

        table = document.add_table(rows=1, cols=5)
        header = table.rows[0].cells
        header[0].text = "ID"
        header[1].text = "检查项"
        header[2].text = "状态"
        header[3].text = "优先级"
        header[4].text = "描述"

        for item in self._active_items(session):
            row = table.add_row().cells
            row[0].text = item.id
            row[1].text = item.content
            row[2].text = item.status
            row[3].text = item.priority
            row[4].text = item.description

        document.save(target)
        return target

    def export_excel(
        self, session: ProjectSession, destination: Path | None = None
    ) -> Path:
        target = destination or (session.root / "report.xlsx")
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Accessibility Report"
        sheet.append(["ID", "内容", "状态", "优先级", "描述", "截图"])
        for item in self._active_items(session):
            sheet.append(
                [
                    item.id,
                    item.content,
                    item.status,
                    item.priority,
                    item.description,
                    item.image_path,
                ]
            )
        workbook.save(target)
        return target

    def export_jira_csv(
        self, session: ProjectSession, destination: Path | None = None
    ) -> Path:
        target = destination or (session.root / "jira_issues.csv")
        with target.open("w", newline="", encoding="utf-8-sig") as handle:
            writer = csv.writer(handle)
            writer.writerow(["Summary", "Description", "Priority", "Issue Type"])
            for item in self._active_items(session):
                if item.status != "failed":
                    continue
                writer.writerow(
                    [
                        f"[A11Y]{item.id} {item.content}",
                        item.description,
                        item.priority.upper(),
                        "Bug",
                    ]
                )
        return target

    def export_markdown(
        self, session: ProjectSession, destination: Path | None = None
    ) -> Path:
        target = destination or (session.root / "report.md")
        lines = [
            "# 信息无障碍测试报告",
            "",
            f"- 项目编号：{session.data.meta.project_number}",
            f"- 名称：{session.data.meta.project_name or '-'}",
            f"- 场景：{session.data.meta.scenario}",
            f"- 模板：{session.data.meta.template}",
            "",
            "| ID | 内容 | 状态 | 优先级 | 描述 | 截图 |",
            "| --- | --- | --- | --- | --- | --- |",
        ]
        for item in self._active_items(session):
            lines.append(
                f"| {item.id} | {item.content} | {item.status} | {item.priority} | "
                f"{item.description or '-'} | {item.image_path or '-'} |"
            )
        target.write_text("\n".join(lines), encoding="utf-8")
        return target

    def _active_items(self, session: ProjectSession) -> list[ProjectItem]:
        return [item for item in session.data.items if not item.deleted]

    def _stats(self, session: ProjectSession) -> dict[str, int]:
        stats = {
            "total": 0,
            "passed": 0,
            "failed": 0,
            "pending": 0,
            "not_applicable": 0,
        }
        for item in self._active_items(session):
            stats["total"] += 1
            key = item.status if item.status in stats else "pending"
            stats[key] += 1
        return stats
